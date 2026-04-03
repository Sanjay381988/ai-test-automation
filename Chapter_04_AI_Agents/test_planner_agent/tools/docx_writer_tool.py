import docx
import os

def generate_test_plan_docx(template_path: str, output_path: str, context: dict) -> bool:
    try:
        # Load the existing template
        doc = docx.Document(template_path)
        
        # Append the new generated test plan information
        doc.add_page_break()
        doc.add_heading('AI Generated Test Plan Details', level=1)
        
        if 'objective' in context:
            doc.add_heading('Objective', level=2)
            doc.add_paragraph(str(context['objective']))
            
        if 'scope' in context:
            doc.add_heading('Scope', level=2)
            doc.add_paragraph(str(context['scope']))
            
        if 'test_scenarios' in context:
            doc.add_heading('Test Scenarios', level=2)
            scenarios = context['test_scenarios']
            if isinstance(scenarios, list):
                for s in scenarios:
                    # Safely handle missing styles in the docx template
                    try:
                        doc.add_paragraph(s, style='List Bullet')
                    except KeyError:
                        doc.add_paragraph(f"- {s}")
            else:
                doc.add_paragraph(str(scenarios))
                
        if 'risks' in context:
            doc.add_heading('Risks', level=2)
            doc.add_paragraph(str(context['risks']))
            
        if 'environment' in context:
            doc.add_heading('Environment', level=2)
            doc.add_paragraph(str(context['environment']))
            
        doc.save(output_path)
        return True
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False
